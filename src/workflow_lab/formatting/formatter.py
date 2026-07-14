#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
GB/T 7714-2015 Citation & Formatting Tool for Scientific Workflow Lab
Codex annotation: Created by Codex on 2026-07-10.
"""

import os
import re
import sys
import argparse
import urllib.parse
from pathlib import Path
import httpx
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- Styles and Constants ---
FONT_ZH = "宋体"
FONT_EN = "Times New Roman"

COLOR_RED = RGBColor(220, 53, 69)      # Used for AI additions/modifications
COLOR_BLACK = RGBColor(0, 0, 0)
COLOR_MUTED = RGBColor(93, 104, 120)

# --- Crossref & Metadata Fetching Helper ---
def fetch_crossref_metadata(query: str) -> dict | None:
    """Fetch metadata from Crossref API using query string."""
    url = f"https://api.crossref.org/works?query={urllib.parse.quote(query)}&rows=1"
    headers = {"User-Agent": "ScientificWorkflowLab/1.0 (mailto:hello@scientificworkflowlab.org)"}
    try:
        response = httpx.get(url, headers=headers, timeout=10.0)
        if response.status_code == 200:
            data = response.json()
            items = data.get("message", {}).get("items", [])
            if items:
                return items[0]
    except Exception as e:
        sys.stderr.write(f"Warning: Crossref lookup failed for query '{query}': {e}\n")
    return None

def format_author_gbt7714(authors_list: list) -> str:
    """Format authors list to GB/T 7714-2015 rule (last name first, up to 3 authors, et al.)."""
    if not authors_list:
        return ""
    
    formatted_authors = []
    for author in authors_list:
        family = author.get("family", "").strip()
        given = author.get("given", "").strip()
        if family:
            # Capitalize family name as per GB/T 7714 (e.g. QUENBY S)
            family_upper = family.upper()
            # Abbreviate given name (first letters of given name parts)
            given_abbr = "".join([part[0].upper() for part in given.split() if part])
            formatted_authors.append(f"{family_upper} {given_abbr}".strip())
        else:
            name = author.get("name", "").strip()
            if name:
                formatted_authors.append(name)
                
    if not formatted_authors:
        return ""
        
    if len(formatted_authors) <= 3:
        return ", ".join(formatted_authors)
    else:
        return ", ".join(formatted_authors[:3]) + ", et al"

def parse_crossref_to_gbt7714(item: dict) -> str:
    """Transform Crossref item metadata to GB/T 7714-2015 reference string."""
    # 1. Authors
    authors_data = item.get("author", [])
    author_str = format_author_gbt7714(authors_data)
    if not author_str:
        author_str = "Anon"
        
    # 2. Title
    title_list = item.get("title", [""])
    title_str = title_list[0].strip() if title_list else ""
    # Remove trailing period if present
    if title_str.endswith("."):
        title_str = title_str[:-1]
        
    # 3. Journal Name
    container_list = item.get("container-title", [""])
    journal_str = container_list[0].strip() if container_list else ""
    
    # 4. Year
    pub_date = item.get("published-print") or item.get("published-online") or item.get("created")
    year_str = ""
    if pub_date:
        parts = pub_date.get("date-parts", [[]])[0]
        if parts:
            year_str = str(parts[0])
            
    # 5. Volume, Issue, Pages
    volume = item.get("volume", "").strip()
    issue = item.get("issue", "").strip()
    
    pages = item.get("page", "").strip()
    # Normalize page delimiters
    if pages:
        pages = pages.replace("--", "-")
        
    # 6. DOI
    doi = item.get("DOI", "").strip()
    
    # Construct Journal Article citation
    citation = f"{author_str}. {title_str} [J]. {journal_str}"
    if year_str:
        citation += f", {year_str}"
    if volume:
        citation += f", {volume}"
    if issue:
        citation += f"({issue})"
    if pages:
        citation += f": {pages}"
    if doi:
        citation += f". DOI: {doi}"
    if not citation.endswith("."):
        citation += "."
        
    return citation

# --- Document Helper Functions ---
def set_run_font(run, name_ascii=FONT_EN, name_h=FONT_ZH, size_pt=10.5, color=None, bold=False, italic=False):
    """Format font styling on a run level for Word output."""
    run.font.name = name_ascii
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name_ascii)
    rFonts.set(qn('w:hAnsi'), name_ascii)
    rFonts.set(qn('w:eastAsia'), name_h)
    rPr.append(rFonts)
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color

def set_paragraph_spacing(paragraph, line_spacing=1.25, space_before=0, space_after=6, keep_together=False):
    """Set line spacing and paragraph margins."""
    p_format = paragraph.paragraph_format
    p_format.line_spacing = line_spacing
    p_format.space_before = Pt(space_before)
    p_format.space_after = Pt(space_after)
    p_format.keep_with_next = keep_together

def configure_page_layout(doc):
    """Apply NSFC standard margins to all sections."""
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

# --- Markdown Parser and Converter ---
class MarkdownToWordConverter:
    def __init__(self, output_path: Path):
        self.doc = Document()
        configure_page_layout(self.doc)
        self.output_path = output_path
        self.references_list = []
        self.in_references = False
        
    def add_heading(self, text: str, level: int, highlight: bool = False):
        """Add custom formatted heading to Word document."""
        p = self.doc.add_paragraph()
        # Heading sizes: H1=15pt (小三), H2=14pt (四号), H3=12pt (小四)
        size_map = {1: 15, 2: 14, 3: 12}
        size = size_map.get(level, 12)
        
        # Heading spacing: larger margin before
        set_paragraph_spacing(p, line_spacing=1.25, space_before=12, space_after=6, keep_together=True)
        
        # Add heading number prefixes if H2/H3 (e.g. 3.1)
        run = p.add_run(text)
        set_run_font(run, name_h=FONT_ZH, size_pt=size, color=COLOR_RED if highlight else COLOR_BLACK, bold=True)
        
    def add_paragraph_with_inline(self, line: str):
        """Parse inline markers like **bold**, *italic*, [red](...) or redline flags."""
        p = self.doc.add_paragraph()
        # NSFC standard: First line indent 2 characters (approx 0.74 cm for 10.5pt font)
        p.paragraph_format.first_line_indent = Cm(0.74)
        set_paragraph_spacing(p, line_spacing=1.25, space_before=0, space_after=6)
        
        # Basic markdown parsing logic for inline formatting
        # We look for redline changes denoted by [Redline: text] or similar
        # For simplicity, we parse bold `**text**` and red text markers like `[red]{text}` or `++text++`
        pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|\+\+.*?\+\+|\[red\]\{.*?\})')
        parts = pattern.split(line)
        
        for part in parts:
            if not part:
                continue
            
            run = p.add_run()
            if part.startswith('**') and part.endswith('**'):
                run.text = part[2:-2]
                set_run_font(run, size_pt=10.5, bold=True)
            elif part.startswith('*') and part.endswith('*'):
                run.text = part[1:-1]
                set_run_font(run, size_pt=10.5, italic=True)
            elif part.startswith('++') and part.endswith('++'):
                # Additions highlighted in RED
                run.text = part[2:-2]
                set_run_font(run, size_pt=10.5, color=COLOR_RED, bold=True)
            elif part.startswith('[red]\{') and part.endswith('\}'):
                # Explicit red text marker
                run.text = part[6:-1]
                set_run_font(run, size_pt=10.5, color=COLOR_RED)
            else:
                run.text = part
                set_run_font(run, size_pt=10.5)

    def add_reference_line(self, ref_line: str, refine: bool = False):
        """Parse, compile, and format single reference entry."""
        # Clean index marker (e.g., "[1]")
        match_idx = re.match(r'^\[(\d+)\]\s*(.*)', ref_line.strip())
        if not match_idx:
            # Fallback if no index matches
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.74)
            p.paragraph_format.first_line_indent = Cm(-0.74)
            set_paragraph_spacing(p, line_spacing=1.0, space_before=0, space_after=3)
            run = p.add_run(ref_line)
            set_run_font(run, size_pt=9.0, color=COLOR_MUTED)
            return

        idx = match_idx.group(1)
        raw_content = match_idx.group(2).strip()
        
        final_citation = raw_content
        
        # If refinement is requested and it contains enough characters to query Crossref
        if refine and len(raw_content) > 15 and not raw_content.startswith("中共中央"):
            sys.stdout.write(f"Querying metadata for reference [{idx}]: {raw_content[:50]}...\n")
            meta = fetch_crossref_metadata(raw_content)
            if meta:
                compiled = parse_crossref_to_gbt7714(meta)
                if compiled:
                    final_citation = compiled
                    sys.stdout.write(f"-> Success: {final_citation}\n")
        
        # Add reference paragraph with Hanging Indent
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.74)
        set_paragraph_spacing(p, line_spacing=1.0, space_before=0, space_after=3)
        
        # Index run
        run_idx = p.add_run(f"[{idx}]\t")
        set_run_font(run_idx, size_pt=9.5, bold=True)
        
        # Citation content run
        run_cit = p.add_run(final_citation)
        set_run_font(run_cit, size_pt=9.5)
        
    def convert_file(self, md_path: Path, refine_citations: bool = False):
        """Process Markdown file line by line and construct Word output."""
        if not md_path.exists():
            raise FileNotFoundError(f"Markdown file {md_path} does not exist.")
            
        with open(md_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            
        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue
                
            # Detect references section
            if re.match(r'^#+\s*(参考文献|References)', line_str):
                self.in_references = True
                self.add_heading("参考文献", level=2)
                continue
                
            if self.in_references:
                # Add to references block
                self.add_reference_line(line_str, refine=refine_citations)
            else:
                # Heading parsing
                match_h = re.match(r'^(#+)\s*(.*)', line_str)
                if match_h:
                    level = len(match_h.group(1))
                    title = match_h.group(2).strip()
                    # Check if heading has custom highlight marker
                    highlight = False
                    if title.startswith("++") and title.endswith("++"):
                        title = title[2:-2]
                        highlight = True
                    self.add_heading(title, level=level, highlight=highlight)
                else:
                    # Normal Paragraph
                    self.add_paragraph_with_inline(line_str)
                    
        self.doc.save(str(self.output_path))
        print(f"\nDocument formatted successfully and saved to: {self.output_path}")

# --- Main Entrypoint ---
def main():
    parser = argparse.ArgumentParser(description="GB/T 7714-2015 Citation & Word Layout Formatter")
    parser.add_argument("-i", "--input", required=True, help="Path to raw Markdown file")
    parser.add_argument("-o", "--output", required=True, help="Path to save formatted .docx file")
    parser.add_argument("-r", "--refine", action="store_true", help="Attempt Crossref API query to auto-fix bibliography")
    
    args = parser.parse_args()
    
    input_path = Path(args.input)
    output_path = Path(args.output)
    
    # Ensure directories exist
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    converter = MarkdownToWordConverter(output_path)
    converter.convert_file(input_path, refine_citations=args.refine)

if __name__ == "__main__":
    main()
